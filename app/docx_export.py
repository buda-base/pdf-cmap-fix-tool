"""Render extracted text to a .docx file (in memory).

Two inputs are supported:
- Markdown (from PyMuPDF4LLM) -> highest fidelity via pandoc when available
  (headings, bold/italic, lists, tables), falling back to a small python-docx
  renderer that covers the common subset.
- Plain text (the legacy-Tibetan conversion path) -> one paragraph per line via
  python-docx, so verse/line structure is preserved verbatim.

python-docx is a hard dependency (always available); pandoc is optional and only
used to improve Markdown fidelity.
"""

from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import tempfile

_PANDOC = shutil.which("pandoc")

# Inline markdown emphasis: **bold**, __bold__, *italic*, _italic_.
_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_)")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
_ORDERED_RE = re.compile(r"^\d+[.)]\s+(.*)$")


def to_docx_bytes(text: str, *, is_markdown: bool, title: str | None = None) -> bytes:
    """Return .docx bytes for `text`. Uses pandoc for Markdown when present."""
    if is_markdown and _PANDOC:
        out = _via_pandoc(text)
        if out is not None:
            return out
    return _via_python_docx(text, is_markdown=is_markdown, title=title)


def _via_pandoc(markdown: str) -> bytes | None:
    """Convert Markdown -> docx with pandoc. Returns None on any failure."""
    fd, out_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        subprocess.run(
            # hard_line_breaks: keep single newlines as line breaks (verse lines).
            ["pandoc", "-f", "markdown+hard_line_breaks", "-t", "docx", "-o", out_path],
            input=markdown.encode("utf-8"),
            check=True,
            capture_output=True,
            timeout=120,
        )
        with open(out_path, "rb") as fh:
            data = fh.read()
        return data or None
    except Exception:
        return None
    finally:
        if os.path.exists(out_path):
            os.remove(out_path)


def _via_python_docx(text: str, *, is_markdown: bool, title: str | None) -> bytes:
    from docx import Document

    doc = Document()
    if title:
        doc.core_properties.title = title

    if not is_markdown:
        # Plain text: preserve every line, blank lines become empty paragraphs.
        for line in text.split("\n"):
            doc.add_paragraph(line)
    else:
        _render_markdown(doc, text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_markdown(doc, markdown: str) -> None:
    """Minimal Markdown -> docx covering the subset PyMuPDF4LLM emits."""
    for raw in markdown.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(heading.group(2).strip(), level=level)
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            _add_inline(doc.add_paragraph(style="List Bullet"), bullet.group(1))
            continue

        ordered = _ORDERED_RE.match(line)
        if ordered:
            _add_inline(doc.add_paragraph(style="List Number"), ordered.group(1))
            continue

        _add_inline(doc.add_paragraph(), line)


def _add_inline(paragraph, text: str) -> None:
    """Add text to a paragraph, honouring **bold** / *italic* markers."""
    for token in _TOKEN_RE.split(text):
        if not token:
            continue
        if (token.startswith("**") and token.endswith("**")) or (
            token.startswith("__") and token.endswith("__")
        ):
            paragraph.add_run(token[2:-2]).bold = True
        elif (token.startswith("*") and token.endswith("*")) or (
            token.startswith("_") and token.endswith("_")
        ):
            paragraph.add_run(token[1:-1]).italic = True
        else:
            paragraph.add_run(token)
